#!/usr/bin/env python

"""
This file opens a docx (Office 2007) file and dumps the text.

If you need to extract text from documents, use this file as a basis for your
work.

Part of Python's docx module - http://github.com/mikemaccana/python-docx
See LICENSE for licensing information.
"""

# This module was taken from a sample of the Python docx module
# Refer its documentation above for details

import sys

from docx import Document

def docxExtract(docxfile):
    try:
        document = Document(docxfile)
        # document = opendocx()
    except:
        print("Error opening docx")
        exit()
    fullText = []
    data = ''
    # Fetch all the text out of the document we just created
    # paratextlist = document.paragraphs
    for para in document.paragraphs:
        fullText.append(para.text)
        data = '\n'.join(fullText)
    # Print out text of document with two newlines under each paragraph
    return data
    

