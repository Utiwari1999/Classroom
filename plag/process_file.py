from docx import Document
from .code.main import process
from .text_matcher.matcher import Matcher,Text
import pdfplumber

def process_file(filename):
    if filename.name.endswith('.docx'):
        document = Document(filename)
        fullText = []
        data = ''
        for para in document.paragraphs:
            fullText.append(para.text)
            data = '\n'.join(fullText)
    else:
        data = ''
        for chunk in filename.chunks():
            data += chunk.decode('utf')
    return process(data)
    
     
def compare_two(filename1,filename2):
    if filename1.name.endswith('.docx'):
        document = Document(filename1)
        fullText = []
        data = ''
        for para in document.paragraphs:
            fullText.append(para.text)
            data = '\n'.join(fullText)
    elif filename1.name.endswith('.pdf'):
        data = ''
        with pdfplumber.open(filename1) as pdf:
            first_page = pdf.pages[0]
            for i in range(len(pdf.pages)):
                data += '\n' + pdf.pages[i].extract_text() 
    else:
        data = ''
        for chunk in filename1.chunks():
            data += chunk.decode('utf')
    
    if filename2.name.endswith('.docx'):
        document = Document(filename2)
        fullText = []
        data2 = ''
        for para in document.paragraphs:
            fullText.append(para.text)
            data2 = '\n'.join(fullText)
    
    elif filename2.name.endswith('.pdf'):
        data2 = ''
        with pdfplumber.open(filename2) as pdf:
            first_page = pdf.pages[0]
            for i in range(len(pdf.pages)):
                data2 += '\n'+ pdf.pages[i].extract_text()
    
    else:
        data2 = ''
        for chunk in filename2.chunks():
            data2 += chunk.decode('utf')
    print(data)
    ta = Text(data, 'A')
    tb = Text(data2, 'B')
    (a,b,c,d) = Matcher(ta,tb).match()
    return a,b,c,d,data,data2
    
     

    